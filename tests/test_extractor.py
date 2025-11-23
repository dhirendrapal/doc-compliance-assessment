from extractor import extract_docx
from docx import Document
import os

def test_extract_docx(tmp_path):
    p = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph('Hello world')
    doc.add_paragraph('Second paragraph')
    doc.save(str(p))
    paras = extract_docx(str(p))
    assert len(paras) == 2
    assert paras[0]['text'] == 'Hello world'