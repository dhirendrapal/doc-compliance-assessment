from fastapi.testclient import TestClient
from docx import Document
from app import app

client = TestClient(app)

def test_modify_after_upload():
    # Step 1: create test docx
    from docx import Document
    doc = Document()
    doc.add_paragraph("She go to market yesterday.")
    doc.save("tests/sample2.docx")

    # Step 2: upload
    with open("tests/sample2.docx", "rb") as f:
        resp = client.post("/upload", files={"file": ("sample2.docx", f)})
    assert resp.status_code == 200

    job_id = resp.json()["job_id"]

    # Step 3: modify using agent
    mod_resp = client.post(f"/modify/{job_id}")
    assert mod_resp.status_code == 200

    # Step 4: download
    dl = client.get(f"/download/{job_id}")
    assert dl.status_code == 200

    with open("out.docx", "wb") as f:
        f.write(dl.content)  # saved to inspect if needed

    doc2 = Document("out.docx")
    assert len(doc2.paragraphs) > 0
